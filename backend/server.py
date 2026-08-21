from dotenv import load_dotenv
from pathlib import Path

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / ".env")

from datetime import datetime, timezone, timedelta
from typing import Any, Dict, List, Optional
import io
import json
import logging
import os
import re
import shutil
import uuid

import certifi
from pymongo import MongoClient
import base64

MONGODB_URI = os.environ.get("MONGODB_URI", "")
mongo_client = MongoClient(MONGODB_URI, tlsCAFile=certifi.where()) if MONGODB_URI else None
db = mongo_client.casefile_db if mongo_client else None


import bcrypt
import jwt
from docx import Document
from docx.shared import Inches
from fastapi import APIRouter, Depends, FastAPI, File, HTTPException, Request, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, ConfigDict, EmailStr, Field
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas
from starlette.middleware.cors import CORSMiddleware

app = FastAPI(title="Casefile Document API", version="2.0.0")
api_router = APIRouter(prefix="/api")

TEMPLATE_DIR = ROOT_DIR / "templates"
CUSTOM_TEMPLATE_DIR = TEMPLATE_DIR / "custom"
CUSTOM_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

JWT_ALGORITHM = "HS256"
JWT_SECRET = os.environ["JWT_SECRET"]
ACCESS_TOKEN_MINUTES = 60 * 12
MAX_FAILED_ATTEMPTS = 5
LOCKOUT_MINUTES = 15
DOC_CACHE_TTL_MINUTES = 30

ADMIN_EMAIL = os.environ["ADMIN_EMAIL"].lower().strip()
ADMIN_NAME = os.environ.get("ADMIN_NAME", "Casefile Admin")
ADMIN_ID = "admin"
ADMIN_PASSWORD_HASH = bcrypt.hashpw(os.environ["ADMIN_PASSWORD"].encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

# In-memory state (resets on restart)
LOGIN_ATTEMPTS: Dict[str, Dict[str, Any]] = {}
DOC_CACHE: Dict[str, Dict[str, Any]] = {}


CANONICAL_FIELDS = {
    "cd_name": ("Corporate debtor name", "Corporate debtor", True),
    "cin": ("CIN", "Corporate debtor", True),
    "nclt_bench": ("NCLT bench", "Corporate debtor", True),
    "cp_ib_number": ("CP (IB) number", "Corporate debtor", True),
    "md_name": ("Managing director name", "Corporate debtor", False),
    "cd_address": ("Corporate debtor address", "Corporate debtor", False),
    "ip_name": ("IP name", "Insolvency professional", True),
    "ibbi_reg_no": ("IBBI registration number", "Insolvency professional", True),
    "afa_validity": ("AFA validity date", "Insolvency professional", False),
    "ip_reg_address": ("Registered address", "Insolvency professional", False),
    "ip_email": ("Registered email", "Insolvency professional", False),
    "process_email": ("Process-specific email", "Insolvency professional", False),
    "cirp_order_date": ("CIRP order date", "CIRP timeline", False),
    "order_upload_date": ("Order upload date", "CIRP timeline", False),
    "pa_date": ("Public announcement date", "CIRP timeline", False),
    "claim_cutoff_date": ("Claim cut-off date", "CIRP timeline", False),
    "loc_date": ("LOC / CoC date", "CIRP timeline", False),
    "loc_filing_date": ("LOC filing date", "Filing specifics", False),
    "loc_ia_number": ("LOC IA number", "Filing specifics", False),
    "nclt_fee": ("NCLT filing fee", "Filing specifics", False),
    "meeting_number": ("Meeting number", "Meeting details", False),
    "meeting_date": ("Meeting date", "Meeting details", False),
    "meeting_time": ("Meeting time", "Meeting details", False),
    "meeting_mode": ("Meeting mode", "Meeting details", False),
    "meeting_venue": ("Meeting venue", "Meeting details", False),
    "notice_date": ("Notice date", "Meeting details", False),
    "evoting_link": ("E-voting link", "Meeting details", False),
    "process_bank": ("Process bank and branch", "Financial details", False),
    "initial_funding": ("Initial funding", "Financial details", False),
    "ip_fee": ("IP fee", "Financial details", False),
    "ip_ope": ("IP out-of-pocket expenses", "Financial details", False),
    "valuer_fee_cap": ("Valuer fee cap", "Financial details", False),
}

TABLE_LABELS = {
    "df_creditors": ("Creditors list", ["Sr. No.", "Financial creditor", "Voting share (%)"]),
    "df_suspended_mgmt": ("Suspended management", ["Sr. No.", "Name", "Designation"]),
    "df_expenses": ("CIRP expenses", ["Sr. No.", "Head of expense", "Amount (INR)"]),
}

TEMPLATE_CONFIG = [
    ("voting-agenda", "Voting Agenda", "CoC meeting", "Voting agenda with creditor and expense schedules.", "Voting_Agenda_Template.docx", ["cd_name", "ip_name", "ibbi_reg_no", "meeting_number", "meeting_date", "meeting_time", "ip_fee", "ip_ope", "valuer_fee_cap", "process_bank", "df_creditors", "df_expenses"]),
    ("constitution-coc", "Constitution of CoC", "CIRP constitution", "Formal constitution notice with the creditor composition.", "Constitution_of_CoC_Template.docx", ["loc_date", "cd_name", "cin", "nclt_bench", "cp_ib_number", "claim_cutoff_date", "df_creditors", "ip_name", "ibbi_reg_no", "afa_validity", "process_email", "ip_email", "ip_reg_address"]),
    ("notice-first-coc", "Notice of 1st CoC", "CoC meeting", "Notice for the first meeting of the Committee of Creditors.", "Notice_1st_CoC_Template.docx", ["cd_name", "cirp_order_date", "meeting_date", "meeting_time", "meeting_mode", "notice_date", "evoting_link", "ip_name", "ibbi_reg_no", "ip_email", "ip_fee", "ip_reg_address", "process_bank", "process_email", "df_creditors", "df_suspended_mgmt"]),
    ("notice-second-coc", "Notice of 2nd CoC", "CoC meeting", "Notice for a subsequent meeting of the Committee of Creditors.", "Notice_2nd_CoC_Template.docx", ["cd_name", "cirp_order_date", "meeting_date", "meeting_time", "meeting_mode", "meeting_venue", "notice_date", "evoting_link", "ip_name", "ibbi_reg_no", "afa_validity", "ip_email", "ip_reg_address", "process_email", "df_creditors", "df_suspended_mgmt"]),
    ("loc-filing", "LOC Filing & CoC Report", "NCLT filing", "Interlocutory application filing with index and list of dates.", "LOC_Filing_Template.docx", ["loc_ia_number", "cp_ib_number", "cd_name", "md_name", "ip_name", "cirp_order_date", "order_upload_date", "pa_date", "claim_cutoff_date", "loc_date", "loc_filing_date", "nclt_fee", "ip_reg_address", "process_email", "ip_email", "ibbi_reg_no", "afa_validity"]),
]


# ---------------- Models ----------------
class LoginInput(BaseModel):
    email: EmailStr
    password: str


class UserOut(BaseModel):
    id: str
    email: str
    name: str
    role: str


class TokenOut(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserOut


class TemplateFieldSpec(BaseModel):
    key: str
    label: str
    section: str = "Details"
    required: bool = False
    placeholder: str = ""


class TemplateTableSpec(BaseModel):
    key: str
    label: str
    columns: List[str]


class CustomTemplateSpec(BaseModel):
    upload_id: str
    name: str
    category: str = "Custom"
    description: str = ""
    fields: List[TemplateFieldSpec]
    table_inputs: List[TemplateTableSpec] = Field(default_factory=list)


class InspectResult(BaseModel):
    upload_id: str
    detected_fields: List[TemplateFieldSpec]
    detected_tables: List[TemplateTableSpec]


class TemplateRecord(BaseModel):
    id: str
    name: str
    category: str
    description: str
    fields: List[dict]
    table_inputs: List[dict]
    source: str = "builtin"


class DocumentInput(BaseModel):
    template_id: str
    values: Dict[str, str] = Field(default_factory=dict)
    tables: Dict[str, List[List[str]]] = Field(default_factory=dict)
    notes: str = ""


class GeneratedDocument(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    template_id: str
    template_name: str
    company_name: str
    status: str
    created_at: str
    values: Dict[str, str]
    tables: Dict[str, List[List[str]]] = Field(default_factory=dict)
    notes: str = ""


# ---------------- Auth ----------------
def create_access_token(user_id: str, email: str) -> str:
    payload = {
        "sub": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(minutes=ACCESS_TOKEN_MINUTES),
        "type": "access",
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


def admin_user() -> Dict[str, str]:
    return {"id": ADMIN_ID, "email": ADMIN_EMAIL, "name": ADMIN_NAME, "role": "admin"}


async def get_current_user(request: Request) -> Dict[str, str]:
    auth_header = request.headers.get("Authorization", "")
    token = auth_header[7:] if auth_header.startswith("Bearer ") else request.cookies.get("access_token")
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Session expired. Please log in again.")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid session token")
    if payload.get("sub") != ADMIN_ID:
        raise HTTPException(status_code=401, detail="Unknown user")
    return admin_user()


# ---------------- Template records ----------------
def template_records():
    records = []
    for ident, name, category, description, filename, keys in TEMPLATE_CONFIG:
        fields = [{"key": key, "label": CANONICAL_FIELDS[key][0], "section": CANONICAL_FIELDS[key][1], "required": CANONICAL_FIELDS[key][2], "placeholder": f"Enter {CANONICAL_FIELDS[key][0].lower()}"} for key in keys if key in CANONICAL_FIELDS]
        tables = [{"key": key, "label": TABLE_LABELS[key][0], "columns": TABLE_LABELS[key][1], "required": False} for key in keys if key in TABLE_LABELS]
        records.append({"id": ident, "name": name, "category": category, "description": description, "fields": fields, "table_inputs": tables, "filename": filename, "field_keys": keys})
    return records


TEMPLATES = template_records()


def default_columns_for_key(key: str) -> List[str]:
    if key in TABLE_LABELS:
        return TABLE_LABELS[key][1]
    return ["Sr. No.", "Description", "Value"]


def humanize_key(key: str) -> str:
    if key in TABLE_LABELS:
        return TABLE_LABELS[key][0]
    stripped = re.sub(r"^df_", "", key)
    return re.sub(r"[_\-]+", " ", stripped).strip().title() or key


def canonical_placeholder(raw: str) -> str:
    key = re.sub(r"[^a-z0-9]", "", raw.lower())
    aliases = {re.sub(r"[^a-z0-9]", "", k.lower()): k for k in list(CANONICAL_FIELDS) + list(TABLE_LABELS)}
    return aliases.get(key, raw)


def iter_paragraphs(container):
    for paragraph in getattr(container, "paragraphs", []):
        yield paragraph
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def extract_placeholders(docx_path) -> List[str]:
    doc = Document(str(docx_path))
    pattern = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")
    seen: List[str] = []
    for container in [doc] + [section.header for section in doc.sections] + [section.footer for section in doc.sections]:
        for paragraph in iter_paragraphs(container):
            for match in pattern.finditer(paragraph.text):
                key = match.group(1).strip()
                if key and key not in seen:
                    seen.append(key)
    return seen


def replace_paragraph(paragraph, values):
    token_pattern = re.compile(r"\{\{\s*([^}]+?)\s*\}\}")
    runs = list(paragraph.runs)
    if not runs:
        if token_pattern.search(paragraph.text):
            paragraph.text = token_pattern.sub(lambda match: str(values.get(canonical_placeholder(match.group(1)), "")), paragraph.text)
        return
    combined = "".join(run.text or "" for run in runs)
    matches = list(token_pattern.finditer(combined))
    if not matches and token_pattern.search(paragraph.text):
        paragraph.text = token_pattern.sub(lambda match: str(values.get(canonical_placeholder(match.group(1)), "")), paragraph.text)
        return
    for match in reversed(matches):
        replacement = str(values.get(canonical_placeholder(match.group(1)), ""))
        start_run = end_run = None
        start_offset = end_offset = 0
        cursor = 0
        for index, run in enumerate(runs):
            run_end = cursor + len(run.text or "")
            if start_run is None and cursor <= match.start() < run_end:
                start_run, start_offset = index, match.start() - cursor
            if cursor < match.end() <= run_end:
                end_run, end_offset = index, match.end() - cursor
                break
            cursor = run_end
        if start_run is None or end_run is None:
            continue
        if start_run == end_run:
            text = runs[start_run].text or ""
            runs[start_run].text = text[:start_offset] + replacement + text[end_offset:]
        else:
            start_text = runs[start_run].text or ""
            end_text = runs[end_run].text or ""
            runs[start_run].text = start_text[:start_offset] + replacement
            for index in range(start_run + 1, end_run):
                runs[index].text = ""
            runs[end_run].text = end_text[end_offset:]
    if token_pattern.search(paragraph.text):
        nodes = [node for node in paragraph._p.iter() if node.tag.endswith("}t")]
        combined_xml = "".join(node.text or "" for node in nodes)
        for match in reversed(list(token_pattern.finditer(combined_xml))):
            first = last = None
            first_offset = last_offset = 0
            cursor = 0
            for index, node in enumerate(nodes):
                end = cursor + len(node.text or "")
                if first is None and cursor <= match.start() < end:
                    first, first_offset = index, match.start() - cursor
                if cursor < match.end() <= end:
                    last, last_offset = index, match.end() - cursor
                    break
                cursor = end
            if first is None or last is None:
                continue
            replacement = str(values.get(canonical_placeholder(match.group(1)), ""))
            if first == last:
                text = nodes[first].text or ""
                nodes[first].text = text[:first_offset] + replacement + text[last_offset:]
            else:
                first_text = nodes[first].text or ""
                last_text = nodes[last].text or ""
                nodes[first].text = first_text[:first_offset] + replacement
                for index in range(first + 1, last):
                    nodes[index].text = ""
                nodes[last].text = last_text[last_offset:]


def add_data_table(paragraph, rows):
    if not rows:
        return
    parent = paragraph._parent
    table = parent.add_table(rows=1, cols=max(len(row) for row in rows), width=Inches(6.2))
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        table.rows[0].cells[index].text = str(value)
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    paragraph._p.addnext(table._tbl)


def build_docx(template, values, tables, notes):
    if template.get("source") == "custom" and template.get("docx_data"):
        doc = Document(io.BytesIO(template["docx_data"]))
    else:
        template_dir = template.get("template_dir", TEMPLATE_DIR)
        doc = Document(str(template_dir / template["filename"]))
    replacements = {key: value for key, value in values.items() if value is not None}
    for container in [doc] + [section.header for section in doc.sections] + [section.footer for section in doc.sections]:
        for paragraph in list(iter_paragraphs(container)):
            for table_key, rows in tables.items():
                if f"{{{{ {table_key} }}}}" in paragraph.text or f"{{{{{table_key}}}}}" in paragraph.text:
                    paragraph.text = re.sub(r"\{\{\s*" + re.escape(table_key) + r"\s*\}\}", "", paragraph.text, flags=re.I)
                    add_data_table(paragraph, rows)
            replace_paragraph(paragraph, replacements)
    if notes.strip():
        doc.add_heading("Additional matter notes", level=2)
        doc.add_paragraph(notes.strip())
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def doc_text(template, values, tables, notes):
    if template.get("source") == "custom" and template.get("docx_data"):
        doc = Document(io.BytesIO(template["docx_data"]))
    else:
        template_dir = template.get("template_dir", TEMPLATE_DIR)
        doc = Document(str(template_dir / template["filename"]))
    chunks = []
    for container in [doc] + [section.header for section in doc.sections] + [section.footer for section in doc.sections]:
        for paragraph in iter_paragraphs(container):
            text = re.sub(r"\{\{\s*([^}]+?)\s*\}\}", lambda m: str(values.get(canonical_placeholder(m.group(1)), "")), paragraph.text)
            if text.strip():
                chunks.append(text.strip())
    for key, rows in tables.items():
        chunks.extend(" | ".join(map(str, row)) for row in rows)
    if notes.strip():
        chunks += ["Additional matter notes", notes.strip()]
    return chunks


# ---------------- Custom template file storage ----------------

def _all_custom_templates() -> List[Dict[str, Any]]:
    if db is not None:
        return list(db.templates.find({}, {"docx_data": 0, "_id": 0}))
    
    records = []
    for meta in CUSTOM_TEMPLATE_DIR.glob("*.json"):
        try:
            records.append(json.loads(meta.read_text(encoding="utf-8")))
        except (OSError, json.JSONDecodeError):
            continue
    records.sort(key=lambda r: r.get("created_at", ""), reverse=True)
    return records

def _load_custom_template(template_id: str) -> Optional[Dict[str, Any]]:
    if db is not None:
        t = db.templates.find_one({"id": template_id}, {"_id": 0})
        return t

    meta = CUSTOM_TEMPLATE_DIR / f"{template_id}.json"
    if not meta.exists():
        return None
    try:
        return json.loads(meta.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def _expand_custom_template(record: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": record["id"],
        "name": record["name"],
        "category": record.get("category", "Custom"),
        "description": record.get("description", ""),
        "fields": record.get("fields", []),
        "table_inputs": record.get("table_inputs", []),
        "filename": record["filename"],
        "source": "custom",
        "field_keys": [f["key"] for f in record.get("fields", [])] + [t["key"] for t in record.get("table_inputs", [])],
    }


def _builtin_public(template: Dict[str, Any]) -> Dict[str, Any]:
    return {key: value for key, value in template.items() if key not in {"filename", "field_keys"}} | {"source": "builtin"}


def list_all_templates() -> List[Dict[str, Any]]:
    return [_builtin_public(t) for t in TEMPLATES] + [
        {key: value for key, value in _expand_custom_template(c).items() if key not in {"filename", "field_keys"}}
        for c in _all_custom_templates()
    ]


def resolve_template(template_id: str) -> Optional[Dict[str, Any]]:
    for template in TEMPLATES:
        if template["id"] == template_id:
            return {**template, "source": "builtin", "template_dir": TEMPLATE_DIR}
    record = _load_custom_template(template_id)
    if not record:
        return None
    expanded = _expand_custom_template(record)
    expanded["template_dir"] = CUSTOM_TEMPLATE_DIR
    return expanded


# ---------------- Doc cache helpers ----------------
def _cache_prune():
    now = datetime.now(timezone.utc)
    for key in [k for k, v in DOC_CACHE.items() if v["expires_at"] < now]:
        DOC_CACHE.pop(key, None)


def _cache_put(record: Dict[str, Any]) -> None:
    _cache_prune()
    DOC_CACHE[record["id"]] = {**record, "expires_at": datetime.now(timezone.utc) + timedelta(minutes=DOC_CACHE_TTL_MINUTES)}


def _cache_get(doc_id: str) -> Optional[Dict[str, Any]]:
    _cache_prune()
    return DOC_CACHE.get(doc_id)


# ---------------- Auth routes ----------------
@api_router.post("/auth/login", response_model=TokenOut)
async def login(payload: LoginInput, request: Request):
    email = payload.email.lower().strip()
    ip = request.client.host if request.client else "unknown"
    identifier = f"{ip}:{email}"
    now = datetime.now(timezone.utc)
    attempt = LOGIN_ATTEMPTS.get(identifier)
    if attempt and attempt.get("locked_until") and attempt["locked_until"] > now:
        raise HTTPException(status_code=429, detail="Too many failed attempts. Try again shortly.")
    email_ok = email == ADMIN_EMAIL
    password_ok = email_ok and bcrypt.checkpw(payload.password.encode("utf-8"), ADMIN_PASSWORD_HASH.encode("utf-8"))
    if not password_ok:
        count = (attempt.get("count", 0) if attempt else 0) + 1
        locked = now + timedelta(minutes=LOCKOUT_MINUTES) if count >= MAX_FAILED_ATTEMPTS else None
        LOGIN_ATTEMPTS[identifier] = {"count": count, "locked_until": locked}
        raise HTTPException(status_code=401, detail="Invalid email or password")
    LOGIN_ATTEMPTS.pop(identifier, None)
    user = admin_user()
    token = create_access_token(user["id"], user["email"])
    return {"access_token": token, "token_type": "bearer", "user": user}


@api_router.post("/auth/logout")
async def logout(current=Depends(get_current_user)):
    return {"ok": True}


@api_router.get("/auth/me", response_model=UserOut)
async def me(current=Depends(get_current_user)):
    return current



# ---------------- Mongo Models ----------------
class MatterInput(BaseModel):
    name: str
    values: Dict[str, Any] = Field(default_factory=dict)
    tables: Dict[str, List[List[str]]] = Field(default_factory=dict)
    timeline: Dict[str, Any] = Field(default_factory=dict)

# ---------------- Mongo Routes ----------------
@api_router.get("/matters")
async def get_matters(current=Depends(get_current_user)):
    if not db: return []
    matters = list(db.matters.find())
    for m in matters: m['_id'] = str(m['_id'])
    return matters

@api_router.post("/matters")
async def create_matter(payload: MatterInput, current=Depends(get_current_user)):
    if not db: raise HTTPException(status_code=500, detail="Database not configured")
    matter = {"id": str(uuid.uuid4()), "name": payload.name, "values": payload.values, "tables": payload.tables, "timeline": payload.timeline, "created_at": datetime.now(timezone.utc).isoformat()}
    db.matters.insert_one(matter)
    matter['_id'] = str(matter['_id'])
    return matter

@api_router.put("/matters/{matter_id}")
async def update_matter(matter_id: str, payload: MatterInput, current=Depends(get_current_user)):
    if not db: raise HTTPException(status_code=500, detail="Database not configured")
    db.matters.update_one({"id": matter_id}, {"$set": {"name": payload.name, "values": payload.values, "tables": payload.tables, "timeline": payload.timeline, "updated_at": datetime.now(timezone.utc).isoformat()}})
    return {"ok": True}

@api_router.delete("/matters/{matter_id}")
async def delete_matter(matter_id: str, current=Depends(get_current_user)):
    if not db: raise HTTPException(status_code=500, detail="Database not configured")
    db.matters.delete_one({"id": matter_id})
    return {"ok": True}

@api_router.get("/profile")
async def get_profile(current=Depends(get_current_user)):
    if not db: return {}
    profile = db.profiles.find_one({"user_id": current["id"]})
    if profile: profile['_id'] = str(profile['_id'])
    return profile or {}

@api_router.put("/profile")
async def update_profile(payload: Dict[str, Any], current=Depends(get_current_user)):
    if not db: raise HTTPException(status_code=500, detail="Database not configured")
    db.profiles.update_one({"user_id": current["id"]}, {"$set": payload}, upsert=True)
    return {"ok": True}

# ---------------- Root ----------------
@api_router.get("/")
async def root():
    return {"message": "Casefile document API ready"}


# ---------------- Template routes ----------------
@api_router.get("/templates", response_model=List[TemplateRecord])
async def get_templates(current=Depends(get_current_user)):
    return list_all_templates()


@api_router.post("/templates/inspect", response_model=InspectResult)
async def inspect_template(file: UploadFile = File(...), current=Depends(get_current_user)):
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    upload_id = str(uuid.uuid4())
    dest = CUSTOM_TEMPLATE_DIR / f"upload-{upload_id}.docx"
    with dest.open("wb") as sink:
        shutil.copyfileobj(file.file, sink)
    try:
        keys = extract_placeholders(dest)
    except Exception as exc:  # noqa: BLE001
        dest.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Could not read DOCX: {exc}")
    if not keys:
        dest.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="No {{placeholders}} were found in this document")
    detected_fields: List[TemplateFieldSpec] = []
    detected_tables: List[TemplateTableSpec] = []
    for key in keys:
        if key.startswith("df_") or key.lower().startswith(("table_", "list_")):
            detected_tables.append(TemplateTableSpec(key=key, label=humanize_key(key), columns=default_columns_for_key(key)))
        else:
            canonical = CANONICAL_FIELDS.get(canonical_placeholder(key))
            if canonical:
                detected_fields.append(TemplateFieldSpec(key=canonical_placeholder(key), label=canonical[0], section=canonical[1], required=canonical[2], placeholder=f"Enter {canonical[0].lower()}"))
            else:
                detected_fields.append(TemplateFieldSpec(key=key, label=humanize_key(key), section="Details", required=False, placeholder=f"Enter {humanize_key(key).lower()}"))
    return InspectResult(upload_id=upload_id, detected_fields=detected_fields, detected_tables=detected_tables)


@api_router.post("/templates", response_model=TemplateRecord)
async def save_custom_template(spec: CustomTemplateSpec, current=Depends(get_current_user)):
    src = CUSTOM_TEMPLATE_DIR / f"upload-{spec.upload_id}.docx"
    if not src.exists():
        raise HTTPException(status_code=404, detail="Upload not found. Please re-upload the template.")
    template_id = f"custom-{uuid.uuid4().hex[:10]}"
    filename = f"{template_id}.docx"
    
    docx_bytes = src.read_bytes()
    
    record = {
        "id": template_id,
        "name": spec.name.strip() or "Custom template",
        "category": spec.category.strip() or "Custom",
        "description": spec.description.strip(),
        "fields": [f.model_dump() for f in spec.fields],
        "table_inputs": [t.model_dump() for t in spec.table_inputs],
        "filename": filename,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    
    if db is not None:
        db.templates.insert_one({**record, "docx_data": docx_bytes})
        src.unlink(missing_ok=True)
    else:
        src.rename(CUSTOM_TEMPLATE_DIR / filename)
        (CUSTOM_TEMPLATE_DIR / f"{template_id}.json").write_text(json.dumps(record, indent=2), encoding="utf-8")
        
    expanded = _expand_custom_template(record)
    return {key: value for key, value in expanded.items() if key not in {"filename", "field_keys", "template_dir", "docx_data"}}



@api_router.delete("/templates/{template_id}")
async def delete_custom_template(template_id: str, current=Depends(get_current_user)):
    if db is not None:
        db.templates.delete_one({"id": template_id})
        return {"ok": True}
        
    record = _load_custom_template(template_id)
    if not record:
        raise HTTPException(status_code=404, detail="Custom template not found")
    (CUSTOM_TEMPLATE_DIR / record["filename"]).unlink(missing_ok=True)
    (CUSTOM_TEMPLATE_DIR / f"{template_id}.json").unlink(missing_ok=True)
    return {"ok": True}



# ---------------- Document generation (cache-only, no DB) ----------------
@api_router.post("/documents", response_model=GeneratedDocument)
async def generate_document(payload: DocumentInput, current=Depends(get_current_user)):
    template = resolve_template(payload.template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    missing_labels: List[str] = []
    for field in template.get("fields", []):
        if field.get("required") and not str(payload.values.get(field["key"], "")).strip():
            missing_labels.append(field.get("label") or field["key"])
    if missing_labels:
        raise HTTPException(status_code=422, detail=f"Required fields missing: {', '.join(missing_labels)}")
    values = {key: str(value).strip() for key, value in payload.values.items()}
    cleaned_tables = {k: [[str(c) for c in row] for row in rows if any(str(c).strip() for c in row)] for k, rows in payload.tables.items() if rows}
    record = {
        "id": str(uuid.uuid4()),
        "template_id": template["id"],
        "template_name": template["name"],
        "company_name": values.get("cd_name") or "Untitled matter",
        "status": "Ready",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "values": values,
        "tables": cleaned_tables,
        "notes": payload.notes,
    }
    _cache_put(record)
    return record


def _render(doc_id: str, file_format: str) -> StreamingResponse:
    item = _cache_get(doc_id)
    if not item or file_format not in {"docx", "pdf"}:
        raise HTTPException(status_code=404, detail="Document not ready. Please generate again.")
    template = resolve_template(item["template_id"])
    if not template:
        raise HTTPException(status_code=404, detail="Template no longer available")
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "-", item["company_name"]).strip("-")[:50] or "casefile-document"
    if file_format == "docx":
        stream = io.BytesIO(build_docx(template, item["values"], item.get("tables", {}), item.get("notes", "")))
        return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{safe_name}-{template["id"]}.docx"'})
    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    y = 800
    pdf.setTitle(template["name"])
    pdf.setFont("Helvetica-Bold", 14)
    for line in doc_text(template, item["values"], item.get("tables", {}), item.get("notes", "")):
        words = line.split()
        current_line = ""
        for word in words:
            candidate = f"{current_line} {word}".strip()
            if stringWidth(candidate, "Helvetica", 9) > 490:
                if y < 55:
                    pdf.showPage()
                    y = 800
                pdf.setFont("Helvetica", 9)
                pdf.drawString(55, y, current_line)
                y -= 14
                current_line = word
            else:
                current_line = candidate
        if y < 55:
            pdf.showPage()
            y = 800
        pdf.setFont("Helvetica-Bold" if y == 800 else "Helvetica", 9)
        pdf.drawString(55, y, current_line)
        y -= 15
    pdf.save()
    stream.seek(0)
    return StreamingResponse(stream, media_type="application/pdf", headers={"Content-Disposition": f'attachment; filename="{safe_name}-{template["id"]}.pdf"'})


@api_router.get("/documents/{document_id}/download/{file_format}")
async def download_document(document_id: str, file_format: str, current=Depends(get_current_user)):
    return _render(document_id, file_format)


@api_router.get("/documents/{document_id}/export/{file_format}")
async def export_document(document_id: str, file_format: str, token: str):
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.PyJWTError:
        raise HTTPException(status_code=401, detail="Invalid token")
    if payload.get("sub") != ADMIN_ID:
        raise HTTPException(status_code=401, detail="Unknown user")
    return _render(document_id, file_format)


app.include_router(api_router)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get("CORS_ORIGINS", "*").split(","),
    allow_methods=["*"],
    allow_headers=["*"],
)
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(name)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)
logger.info("Casefile API ready. Admin: %s", ADMIN_EMAIL)
