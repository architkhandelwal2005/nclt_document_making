"""Tests for new features: Template Uploader + Case Drafts."""
import io
import os
import shutil
import uuid
import zipfile
from pathlib import Path

import pytest
import requests
from docx import Document

if not os.environ.get("REACT_APP_BACKEND_URL"):
    env_path = Path(__file__).resolve().parents[2] / "frontend" / ".env"
    for line in env_path.read_text().splitlines():
        if line.startswith("REACT_APP_BACKEND_URL="):
            os.environ["REACT_APP_BACKEND_URL"] = line.split("=", 1)[1].strip().strip('"')
            break

BASE_URL = os.environ["REACT_APP_BACKEND_URL"].rstrip("/")
API = f"{BASE_URL}/api"

ADMIN_EMAIL = "arun@casefile.app"
ADMIN_PASSWORD = "Casefile2026!"


@pytest.fixture(scope="module")
def token():
    r = requests.post(f"{API}/auth/login", json={"email": ADMIN_EMAIL, "password": ADMIN_PASSWORD})
    assert r.status_code == 200, r.text
    return r.json()["access_token"]


@pytest.fixture(scope="module")
def hdr(token):
    return {"Authorization": f"Bearer {token}"}


def _make_docx_with_placeholders(placeholders):
    doc = Document()
    doc.add_heading("TEST_ Custom template", level=1)
    for p in placeholders:
        doc.add_paragraph(f"Field {p}: {{{{ {p} }}}}")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _make_docx_no_placeholders():
    doc = Document()
    doc.add_paragraph("Just some plain text, no tokens here.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------- Auth guards ----------------
class TestAuthGuards:
    @pytest.mark.parametrize("method,path", [
        ("get", "/templates"),
        ("post", "/templates/inspect"),
        ("post", "/templates"),
        ("delete", "/templates/foo"),
        ("get", "/drafts"),
        ("get", "/drafts/foo"),
        ("post", "/drafts"),
        ("delete", "/drafts/foo"),
    ])
    def test_endpoints_require_auth(self, method, path):
        r = requests.request(method, f"{API}{path}")
        assert r.status_code == 401, f"{method} {path} -> {r.status_code}"


# ---------------- Template Uploader ----------------
class TestTemplateUploader:
    def test_inspect_rejects_non_docx(self, hdr):
        files = {"file": ("notes.txt", b"hello", "text/plain")}
        r = requests.post(f"{API}/templates/inspect", files=files, headers=hdr)
        assert r.status_code == 400
        assert "docx" in r.json()["detail"].lower()

    def test_inspect_rejects_docx_with_no_placeholders(self, hdr):
        buf = _make_docx_no_placeholders()
        files = {"file": ("empty.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{API}/templates/inspect", files=files, headers=hdr)
        assert r.status_code == 400
        assert "placeholder" in r.json()["detail"].lower()

    def test_inspect_returns_fields_and_tables(self, hdr):
        buf = _make_docx_with_placeholders(["cd_name", "custom_field_xyz", "df_creditors", "df_custom_list"])
        files = {"file": ("t.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{API}/templates/inspect", files=files, headers=hdr)
        assert r.status_code == 200, r.text
        data = r.json()
        assert "upload_id" in data and data["upload_id"]
        field_keys = {f["key"] for f in data["detected_fields"]}
        table_keys = {t["key"] for t in data["detected_tables"]}
        assert "cd_name" in field_keys
        assert "custom_field_xyz" in field_keys
        assert "df_creditors" in table_keys
        assert "df_custom_list" in table_keys
        # df_creditors has known columns
        creditors = next(t for t in data["detected_tables"] if t["key"] == "df_creditors")
        assert creditors["columns"] == ["Sr. No.", "Financial creditor", "Voting share (%)"]
        # Unknown table gets default columns
        custom_list = next(t for t in data["detected_tables"] if t["key"] == "df_custom_list")
        assert custom_list["columns"] and len(custom_list["columns"]) >= 2

    def test_full_custom_template_lifecycle(self, hdr, token):
        # 1. Inspect
        buf = _make_docx_with_placeholders(["cd_name", "df_creditors"])
        files = {"file": ("lc.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = requests.post(f"{API}/templates/inspect", files=files, headers=hdr)
        assert r.status_code == 200
        insp = r.json()
        upload_id = insp["upload_id"]

        # 2. Save custom template
        spec = {
            "upload_id": upload_id,
            "name": "TEST_Custom LC Template",
            "category": "TEST_Custom",
            "description": "e2e test",
            "fields": insp["detected_fields"],
            "table_inputs": insp["detected_tables"],
        }
        r2 = requests.post(f"{API}/templates", json=spec, headers=hdr)
        assert r2.status_code == 200, r2.text
        rec = r2.json()
        assert rec["source"] == "custom"
        assert rec["name"] == "TEST_Custom LC Template"
        assert rec["id"].startswith("custom-")
        template_id = rec["id"]

        # 3. GET /templates includes builtins + this custom one
        r3 = requests.get(f"{API}/templates", headers=hdr)
        assert r3.status_code == 200
        all_templates = r3.json()
        builtin_ids = {"voting-agenda", "constitution-coc", "notice-first-coc", "notice-second-coc", "loc-filing"}
        got_ids = {t["id"] for t in all_templates}
        assert builtin_ids <= got_ids
        assert template_id in got_ids
        found = next(t for t in all_templates if t["id"] == template_id)
        assert found["source"] == "custom"
        # builtins report source=builtin
        for t in all_templates:
            if t["id"] in builtin_ids:
                assert t["source"] == "builtin"

        # 4. Generate a document from the custom template
        gen_payload = {
            "template_id": template_id,
            "values": {"cd_name": "TEST_Custom Co"},
            "tables": {"df_creditors": [["Sr. No.", "Financial creditor", "Voting share (%)"], ["1", "TEST_X Bank", "100"]]},
            "notes": "",
        }
        r4 = requests.post(f"{API}/documents", json=gen_payload, headers=hdr)
        assert r4.status_code == 200, r4.text
        doc_id = r4.json()["id"]

        # 5. Download DOCX + verify replacement + table
        r5 = requests.get(f"{API}/documents/{doc_id}/export/docx", params={"token": token})
        assert r5.status_code == 200
        assert len(r5.content) > 5000
        zf = zipfile.ZipFile(io.BytesIO(r5.content))
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        assert "TEST_Custom Co" in xml
        assert "{{cd_name}}" not in xml
        assert "TEST_X Bank" in xml

        # 6. PDF download
        r6 = requests.get(f"{API}/documents/{doc_id}/export/pdf", params={"token": token})
        assert r6.status_code == 200
        assert r6.content.startswith(b"%PDF")

        # 7. DELETE builtin returns 404
        r7 = requests.delete(f"{API}/templates/constitution-coc", headers=hdr)
        assert r7.status_code == 404

        # 8. DELETE the custom template
        r8 = requests.delete(f"{API}/templates/{template_id}", headers=hdr)
        assert r8.status_code == 200
        # Now GET should not include it
        r9 = requests.get(f"{API}/templates", headers=hdr)
        assert template_id not in {t["id"] for t in r9.json()}
        # Delete again -> 404
        r10 = requests.delete(f"{API}/templates/{template_id}", headers=hdr)
        assert r10.status_code == 404


# ---------------- Drafts ----------------
class TestDrafts:
    def test_draft_crud_and_upsert(self, hdr):
        # Create
        payload = {
            "name": "TEST_Draft One",
            "template_id": "notice-first-coc",
            "values": {"cd_name": "TEST_ACME", "meeting_date": "2025-03-01"},
            "tables": {"df_creditors": [["Sr. No.", "Financial creditor", "Voting share (%)"], ["1", "TEST_FC", "100"]]},
            "notes": "draft notes",
        }
        r = requests.post(f"{API}/drafts", json=payload, headers=hdr)
        assert r.status_code == 200, r.text
        d = r.json()
        assert d["id"] and d["template_name"] == "Notice of 1st CoC"
        assert d["values"]["cd_name"] == "TEST_ACME"
        draft_id = d["id"]

        # Update via upsert with same id
        payload2 = {**payload, "id": draft_id, "name": "TEST_Draft One (updated)", "values": {"cd_name": "TEST_ACME 2"}}
        r2 = requests.post(f"{API}/drafts", json=payload2, headers=hdr)
        assert r2.status_code == 200
        assert r2.json()["id"] == draft_id
        assert r2.json()["name"] == "TEST_Draft One (updated)"

        # GET single
        r3 = requests.get(f"{API}/drafts/{draft_id}", headers=hdr)
        assert r3.status_code == 200
        assert r3.json()["values"]["cd_name"] == "TEST_ACME 2"

        # GET list — draft should be there, sorted desc by updated_at
        r4 = requests.get(f"{API}/drafts", headers=hdr)
        assert r4.status_code == 200
        drafts = r4.json()
        ids = [d["id"] for d in drafts]
        assert draft_id in ids
        # First entry should be most recently updated — our draft (just updated)
        assert drafts[0]["id"] == draft_id

        # Nonexistent draft -> 404
        r5 = requests.get(f"{API}/drafts/{uuid.uuid4()}", headers=hdr)
        assert r5.status_code == 404

        # DELETE
        r6 = requests.delete(f"{API}/drafts/{draft_id}", headers=hdr)
        assert r6.status_code == 200
        # Second delete -> 404
        r7 = requests.delete(f"{API}/drafts/{draft_id}", headers=hdr)
        assert r7.status_code == 404
        # GET after delete -> 404
        r8 = requests.get(f"{API}/drafts/{draft_id}", headers=hdr)
        assert r8.status_code == 404

    def test_draft_rejects_bad_template(self, hdr):
        r = requests.post(f"{API}/drafts", json={"name": "TEST_bad", "template_id": "nope-nope"}, headers=hdr)
        assert r.status_code == 404
